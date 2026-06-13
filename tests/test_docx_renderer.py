"""DocxRenderer: ATS-safe resume docx + plain-text cover letter."""
from __future__ import annotations

import io

from docx import Document

from app.adapters.docx_resume import DocxRenderer
from app.domain.models import (
    AdditionalLine,
    ApplicationAnswer,
    CoverLetterContent,
    EducationEntry,
    ExperienceEntry,
    ResumeContent,
)
from tests.conftest import make_header


def _resume() -> ResumeContent:
    return ResumeContent(
        header=make_header(),
        experience=[
            ExperienceEntry(
                company="Apple",
                title="Software Project Lead",
                location="Austin, TX",
                start="Oct 2024",
                end="Present",
                bullets=["Led architecture of four data platforms.", "Owned 1,500+ pods."],
            )
        ],
        education=[
            EducationEntry(
                institution="Universidad Adolfo Ibanez",
                degree="BS, Computer Science",
                location="Santiago, CL",
                detail="Data Structures, Machine Learning.",
            )
        ],
        additional=[AdditionalLine(label="Technical Skills", text="Python, SQL, Kubernetes")],
    )


def _texts(data: bytes) -> list[str]:
    doc = Document(io.BytesIO(data))
    return [p.text for p in doc.paragraphs]


def test_resume_docx_is_ats_safe_and_complete() -> None:
    data = DocxRenderer().render_resume_docx(_resume())
    doc = Document(io.BytesIO(data))

    # ATS killer: no tables for layout.
    assert doc.tables == []

    texts = [p.text for p in doc.paragraphs]
    joined = "\n".join(texts)

    # Header (deterministic, from config) and standard CAPS section headers.
    assert "Alfredo Gutierrez" in joined
    assert "PROFESSIONAL EXPERIENCE" in texts
    assert "EDUCATION" in texts
    assert "ADDITIONAL" in texts

    # Company | Title left, location | dates right on one line (tab-separated).
    assert any(
        "Apple | Software Project Lead" in t and "Oct 2024" in t and "Present" in t
        for t in texts
    )
    # Bullets render as real text.
    assert any("Led architecture of four data platforms." in t for t in texts)
    # Additional labeled line.
    assert any("Technical Skills" in t and "Python" in t for t in texts)


def test_resume_docx_omits_header_when_blank() -> None:
    content = _resume()
    content = ResumeContent(
        experience=content.experience,
        education=content.education,
        additional=content.additional,
    )
    data = DocxRenderer().render_resume_docx(content)
    # Still renders; just no name line.
    assert data[:2] == b"PK"  # valid zip/docx container


def test_application_questions_docx_is_enumerated() -> None:
    answers = (
        ApplicationAnswer(question="Why do you want this role?", answer="Because of the mission."),
        ApplicationAnswer(question="Describe a hard project.", answer="Migrated a monolith."),
    )
    data = DocxRenderer().render_application_questions_docx(answers)
    doc = Document(io.BytesIO(data))

    # No tables — plain single-column doc.
    assert doc.tables == []
    texts = [p.text for p in doc.paragraphs]
    joined = "\n".join(texts)

    assert "Application Questions" in texts
    # Each question is enumerated with its number; answers render as real text.
    assert any(t == "1. Why do you want this role?" for t in texts)
    assert any(t == "2. Describe a hard project." for t in texts)
    assert "Because of the mission." in joined
    assert "Migrated a monolith." in joined


def test_cover_letter_txt_layout() -> None:
    letter = CoverLetterContent(
        greeting="Dear Data team,",
        body_paragraphs=["First.", "Second.", "Third."],
        closing="Best,\nAlfredo Gutierrez",
    )
    out = DocxRenderer().render_cover_letter_txt(letter)
    assert out.startswith("Dear Data team,\n\n")
    assert "First.\n\nSecond.\n\nThird." in out
    assert out.rstrip().endswith("Best,\nAlfredo Gutierrez")

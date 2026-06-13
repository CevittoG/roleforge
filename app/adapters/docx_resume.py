"""DocumentRenderer adapter.

Layout lives here, never in model output:
  - resume  -> .docx via python-docx (uploaded to Drive as a Google Doc)
  - cover letter -> plain text
  - match report -> Markdown via Jinja

The resume is deliberately ATS- and AI-screener-safe: single column, real
selectable text, standard CAPS section headers, and right-aligned
location/dates via a right TAB STOP (never a table — tables break ATS parsers).
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from jinja2 import Environment, FileSystemLoader, select_autoescape

from app.domain.models import (
    ApplicationAnswer,
    AuditFields,
    CoverLetterContent,
    ResumeContent,
)

_TEMPLATES = Path(__file__).resolve().parent.parent / "templates"

_BODY_FONT = "Garamond"
_BODY_PT = 10.5
_PAGE_W_IN = 8.5      # US Letter
_PAGE_H_IN = 11.0
_MARGIN_IN = 0.8


class DocxRenderer:
    def __init__(self) -> None:
        self._env = Environment(
            loader=FileSystemLoader(_TEMPLATES),
            autoescape=select_autoescape(["html", "xml"]),
        )

    def render_resume_docx(self, content: ResumeContent) -> bytes:
        doc = Document()
        for section in doc.sections:
            section.page_width = Inches(_PAGE_W_IN)
            section.page_height = Inches(_PAGE_H_IN)
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(_MARGIN_IN)
            section.right_margin = Inches(_MARGIN_IN)
        usable = Inches(_PAGE_W_IN - 2 * _MARGIN_IN)  # right tab-stop for dates/location

        normal = doc.styles["Normal"]
        normal.font.name = _BODY_FONT
        normal.font.size = Pt(_BODY_PT)
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.line_spacing = 1.04

        _header(doc, content)

        if content.summary:
            _section_heading(doc, "Summary")
            doc.add_paragraph(content.summary)

        if content.experience:
            _section_heading(doc, "Professional Experience")
            for e in content.experience:
                left = " | ".join(x for x in (e.company, e.title) if x)
                right = " | ".join(x for x in (e.location, _dates(e.start, e.end)) if x)
                _two_col_line(doc, usable, left, right)
                for bullet in e.bullets:
                    p = doc.add_paragraph(_clean(bullet), style="List Bullet")
                    p.paragraph_format.space_after = Pt(0)

        if content.education:
            _section_heading(doc, "Education")
            for ed in content.education:
                left = " | ".join(x for x in (ed.institution, ed.degree) if x)
                right = " | ".join(x for x in (ed.location, ed.dates) if x)
                _two_col_line(doc, usable, left, right)
                if ed.detail:
                    doc.add_paragraph(ed.detail)

        if content.additional:
            _section_heading(doc, "Additional")
            for line in content.additional:
                p = doc.add_paragraph()
                label = p.add_run(f"{line.label}: ")
                label.bold = True
                p.add_run(line.text)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def render_cover_letter_txt(self, content: CoverLetterContent) -> str:
        blocks: list[str] = [content.greeting.strip(), ""]
        for para in content.body_paragraphs:
            blocks.append(para.strip())
            blocks.append("")
        blocks.append(content.closing.strip())
        return "\n".join(blocks).strip() + "\n"

    def render_match_report(self, audit: AuditFields) -> str:
        return self._env.get_template("match_report.md.j2").render(audit=audit)

    def render_application_questions_docx(
        self, answers: tuple[ApplicationAnswer, ...]
    ) -> bytes:
        """An editable Word file: a title, then each question enumerated and bold
        with its answer below. Plain single column — easy to copy into a form."""
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(_MARGIN_IN)
            section.right_margin = Inches(_MARGIN_IN)

        normal = doc.styles["Normal"]
        normal.font.name = _BODY_FONT
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)

        heading = doc.add_paragraph()
        title = heading.add_run("Application Questions")
        title.bold = True
        title.font.size = Pt(16)

        for i, a in enumerate(answers, start=1):
            q = doc.add_paragraph()
            q.paragraph_format.space_before = Pt(12)
            q.paragraph_format.space_after = Pt(2)
            q_run = q.add_run(f"{i}. {a.question}")
            q_run.bold = True
            doc.add_paragraph(_clean(a.answer))

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()


# --- docx helpers (python-docx is untyped; annotate at the boundary) ---


def _header(doc: Any, content: ResumeContent) -> None:
    h = content.header
    if h.name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h.name)
        run.bold = True
        run.font.size = Pt(18)
    bits = [b for b in (h.location, h.email, h.phone, *h.links) if b]
    if bits:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("  ·  ".join(bits))
        run.font.size = Pt(9.5)


def _section_heading(doc: Any, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    _bottom_border(p)


def _two_col_line(doc: Any, usable_width: Any, left: str, right: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    if right:
        p.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
    left_run = p.add_run(left)
    left_run.bold = True
    if right:
        p.add_run("\t")
        p.add_run(right)


def _bottom_border(paragraph: Any) -> None:
    """Add a single bottom rule under a paragraph (no direct python-docx API)."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    p_pr.append(borders)


def _dates(start: str, end: str) -> str:
    start, end = start.strip(), end.strip()
    if start and end:
        return f"{start} – {end}"  # noqa: RUF001 — en dash is the correct date-range glyph
    return start or end


def _clean(text: str) -> str:
    """Bullets are plain text; strip any stray markdown bold the model emitted."""
    return text.strip().replace("**", "")
